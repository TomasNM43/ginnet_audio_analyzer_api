import os
import io
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import tempfile
from typing import List, Dict, Optional

# ReportLab para PDFs
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY


# ─── Informe de Autenticidad de Video (YOLO) ─────────────────────────────────

def generate_yolo_report(all_detections: List[Dict], device: str = 'cpu') -> bytes:
    """
    Genera el informe DOCX de detección YOLOv8.
    Retorna el documento como bytes.
    """
    doc = Document()

    title = doc.add_heading('Informe de Detección YOLOv8 en Video', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Fecha de análisis: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph()

    doc.add_heading('Resumen Ejecutivo', level=1)
    total_videos = len(all_detections)
    total_detections = sum(len(v['detections']) for v in all_detections)

    summary = doc.add_paragraph()
    summary.add_run('Total de videos analizados: ').bold = True
    summary.add_run(f'{total_videos}\n')
    summary.add_run('Total de objetos detectados: ').bold = True
    summary.add_run(f'{total_detections}\n')
    summary.add_run('Modelo YOLO usado: ').bold = True
    if all_detections and 'modelo_usado' in all_detections[0]:
        summary.add_run(f'{all_detections[0]["modelo_usado"]}\n')
    else:
        summary.add_run('No especificado\n')

    if total_detections > 0:
        all_conf = [d['confidence'] for v in all_detections for d in v['detections']]
        avg_conf = float(np.mean(all_conf))
        avg_area = sum(d['area_pixels'] for v in all_detections for d in v['detections']) / total_detections
        unique_classes = set(d['class_name'] for v in all_detections for d in v['detections'])

        summary.add_run('Estado: ').bold = True
        summary.add_run('Se detectaron objetos en el/los video(s)\n')
        summary.add_run('Confianza promedio: ').bold = True
        summary.add_run(f'{avg_conf:.3f}\n')
        summary.add_run('Clases detectadas: ').bold = True
        summary.add_run(f'{", ".join(unique_classes)}\n')
        summary.add_run('Área promedio detectada: ').bold = True
        summary.add_run(f'{avg_area:.0f} píxeles²')
    else:
        summary.add_run('Estado: ').bold = True
        summary.add_run('No se detectaron objetos significativos')

    if total_detections > 0:
        doc.add_page_break()
        doc.add_heading('Análisis Detallado por Video', level=1)

        for i, video_data in enumerate(all_detections, 1):
            if not video_data['detections']:
                continue

            doc.add_heading(f'Video {i}: {video_data["video_name"]}', level=2)

            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            headers = ['Frame', 'Tiempo (s)', 'Clase', 'Confianza', 'X', 'Y', 'Ancho', 'Alto', 'Área (%)']
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
                table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

            for det in video_data['detections']:
                x, y, w, h = det['bbox']
                row = table.add_row().cells
                row[0].text = str(det['frame'])
                row[1].text = f"{det['time']:.2f}"
                row[2].text = det['class_name']
                row[3].text = f"{det['confidence']:.3f}"
                row[4].text = str(x)
                row[5].text = str(y)
                row[6].text = str(w)
                row[7].text = str(h)
                row[8].text = f"{det['area_percentage']:.2f}%"

            doc.add_paragraph()

    doc.add_heading('Metodología y Tecnología', level=1)
    meth = doc.add_paragraph()
    meth.add_run('Técnica de detección: ').bold = True
    meth.add_run('YOLOv8 (You Only Look Once versión 8)\n')
    meth.add_run('Procesamiento: ').bold = True
    meth.add_run('Detección de objetos en tiempo real con bounding boxes y confianza\n')
    meth.add_run('Umbral de confianza: ').bold = True
    meth.add_run('0.5 (50%)\n')
    meth.add_run('Dispositivo de procesamiento: ').bold = True
    meth.add_run(device)

    return _doc_to_bytes(doc)


# ─── Informe de Continuidad ───────────────────────────────────────────────────

def generate_continuity_report(analysis_results: List[Dict]) -> bytes:
    """
    Genera el informe DOCX de análisis de continuidad.
    Retorna el documento como bytes.
    """
    doc = Document()

    title = doc.add_heading('Informe de Análisis de Continuidad de Video', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Fecha de análisis: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph()

    doc.add_heading('Resumen Ejecutivo', level=1)
    total_videos = len(analysis_results)
    total_disc = sum(len(r['discontinuities']) for r in analysis_results)

    summary = doc.add_paragraph()
    summary.add_run('Total de videos analizados: ').bold = True
    summary.add_run(f'{total_videos}\n')
    summary.add_run('Total de discontinuidades detectadas: ').bold = True
    summary.add_run(f'{total_disc}\n')

    if total_disc > 0:
        summary.add_run('Estado: ').bold = True
        summary.add_run('Se detectaron posibles cortes o ediciones\n')
        summary.add_run('Promedio por video: ').bold = True
        summary.add_run(f'{total_disc / total_videos:.1f}\n')
    else:
        summary.add_run('Estado: ').bold = True
        summary.add_run('No se detectaron discontinuidades significativas')

    doc.add_page_break()

    for i, result in enumerate(analysis_results, 1):
        doc.add_heading(f'Análisis Detallado - Video {i}', level=1)

        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_data = [
            ('Nombre del archivo:', result['video_name']),
            ('Total de frames:', str(result['total_frames'])),
            ('FPS:', f"{result['fps']:.2f}"),
            ('Duración:', f"{result['duration']:.2f} segundos"),
            ('Distancia máxima:', f"{result['max_distance']:.4f}"),
            ('Discontinuidades detectadas:', str(len(result['discontinuities'])))
        ]
        for row, (label, value) in enumerate(info_data):
            info_table.cell(row, 0).text = label
            info_table.cell(row, 0).paragraphs[0].runs[0].bold = True
            info_table.cell(row, 1).text = value

        doc.add_paragraph()

        # Insertar gráfico si viene en base64
        if result.get('plot_base64'):
            doc.add_heading('Gráfico de Análisis de Continuidad', level=2)
            try:
                plot_bytes = base64.b64decode(result['plot_base64'])
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(plot_bytes)
                    tmp_path = tmp.name
                doc.add_picture(tmp_path, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                os.unlink(tmp_path)
            except Exception as e:
                doc.add_paragraph(f'Error al insertar gráfico: {e}')

        if result['discontinuities']:
            doc.add_heading('Discontinuidades Detectadas', level=2)
            disc_table = doc.add_table(rows=1, cols=4)
            disc_table.style = 'Table Grid'
            for j, h in enumerate(['Frame', 'Tiempo (s)', 'Tiempo (mm:ss)', 'Distancia']):
                disc_table.rows[0].cells[j].text = h
                disc_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

            for disc in result['discontinuities']:
                row = disc_table.add_row().cells
                row[0].text = str(disc['frame'])
                row[1].text = f"{disc['time']:.2f}"
                row[2].text = disc.get('time_formatted', '')
                row[3].text = f"{disc['distance']:.4f}"
        else:
            doc.add_paragraph('No se detectaron discontinuidades significativas.')

        if i < len(analysis_results):
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading('Metodología Empleada', level=1)
    meth = doc.add_paragraph()
    meth.add_run('Técnica utilizada: ').bold = True
    meth.add_run('Análisis de correlación de histogramas\n')
    meth.add_run('Umbral de detección: ').bold = True
    meth.add_run('Distancia > 1.0\n')
    meth.add_run('Procesamiento: ').bold = True
    meth.add_run('Conversión a escala de grises, histogramas normalizados, '
                 'comparación por correlación entre frames consecutivos.')

    return _doc_to_bytes(doc)


# ─── Informe de Extracción de Fotogramas ─────────────────────────────────────

def generate_frame_extraction_report(video_name: str, config: Dict,
                                     extraction_info: List[Dict]) -> bytes:
    """
    Genera el informe DOCX de extracción de fotogramas.
    Retorna el documento como bytes.
    """
    doc = Document()

    title = doc.add_heading('Informe de Extracción de Fotogramas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph()

    doc.add_heading('Configuración de Extracción', level=1)
    cfg_table = doc.add_table(rows=7, cols=2)
    cfg_table.style = 'Table Grid'
    cfg_data = [
        ('Video:', video_name),
        ('Frame inicial:', str(config.get('start_frame', 1))),
        ('Frame final:', str(config.get('end_frame', '-'))),
        ('Salto entre frames:', str(config.get('skip_frames', 1))),
        ('Ajuste de brillo:', f"{config.get('brightness_adjustment', 20):+d}"),
        ('Guardar en color:', 'Sí' if config.get('color_frames') else 'No'),
        ('Guardar en escala de grises:', 'Sí' if config.get('grayscale_frames') else 'No'),
    ]
    for row, (label, value) in enumerate(cfg_data):
        cfg_table.cell(row, 0).text = label
        cfg_table.cell(row, 0).paragraphs[0].runs[0].bold = True
        cfg_table.cell(row, 1).text = value

    doc.add_paragraph()
    doc.add_heading(f'Fotogramas Extraídos ({len(extraction_info)})', level=1)

    if extraction_info:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        for j, h in enumerate(['Frame #', 'Tiempo (s)', 'Tiempo (mm:ss)', 'Brillo aplicado']):
            tbl.rows[0].cells[j].text = h
            tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True

        for info in extraction_info:
            row = tbl.add_row().cells
            row[0].text = str(info['frame_number'])
            row[1].text = f"{info['time_seconds']:.2f}"
            row[2].text = info.get('time_formatted', '')
            row[3].text = f"{info['brightness_applied']:+d}"

    return _doc_to_bytes(doc)


# ─── Informe de Conversión a Escala de Grises ─────────────────────────────────

def generate_grayscale_report(conversion_result: Dict) -> bytes:
    """
    Genera el informe DOCX de conversión a escala de grises.
    Retorna el documento como bytes.
    """
    doc = Document()

    title = doc.add_heading('Informe de Conversión a Escala de Grises', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph()

    doc.add_heading('Resumen de Conversión', level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    total = conversion_result.get('total_files', 0)
    success = conversion_result.get('converted_count', 0)
    error_count = conversion_result.get('error_count', 0)
    rate = (success / total * 100) if total > 0 else 0

    for row, (label, value) in enumerate([
        ('Total de archivos:', str(total)),
        ('Conversiones exitosas:', str(success)),
        ('Errores:', str(error_count)),
        ('Tasa de éxito:', f'{rate:.1f}%'),
    ]):
        summary_table.cell(row, 0).text = label
        summary_table.cell(row, 0).paragraphs[0].runs[0].bold = True
        summary_table.cell(row, 1).text = value

    doc.add_paragraph()

    if conversion_result.get('converted'):
        doc.add_heading('Imágenes Convertidas', level=1)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        for j, h in enumerate(['Original', 'Convertido', 'Resolución', 'Tamaño orig. (KB)', 'Reducción (%)']):
            tbl.rows[0].cells[j].text = h
            tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True

        for img in conversion_result['converted']:
            row = tbl.add_row().cells
            row[0].text = img['original_name']
            row[1].text = img['output_name']
            row[2].text = f"{img['original_size'][0]}x{img['original_size'][1]}"
            row[3].text = f"{img['original_file_size_kb']:.1f}"
            row[4].text = f"{img['size_reduction_pct']:.1f}%"

    if conversion_result.get('errors'):
        doc.add_paragraph()
        doc.add_heading('Errores', level=1)
        for err in conversion_result['errors']:
            doc.add_paragraph(f'• {err}')

    doc.add_paragraph()
    doc.add_heading('Información Técnica', level=1)
    tech = doc.add_paragraph()
    tech.add_run('Método: ').bold = True
    tech.add_run('OpenCV cv2.cvtColor() con CV_BGR2GRAY\n')
    tech.add_run('Formato de salida: ').bold = True
    tech.add_run('PNG para preservar calidad')

    return _doc_to_bytes(doc)


# ─── Helper ───────────────────────────────────────────────────────────────────

def _doc_to_bytes(doc: Document) -> bytes:
    """Serializa un Document de python-docx a bytes."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# ─── Informe de Autenticidad de Video (YOLO) - PDF ──────────────────────────────────────────────

def generate_yolo_report_pdf(all_detections: List[Dict], device: str = 'cpu') -> bytes:
    """
    Genera el informe PDF de detección YOLOv8.
    Retorna el documento como bytes.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    story = []
    styles = getSampleStyleSheet()
    
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a5490'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#1a5490'),
        spaceAfter=12,
        spaceBefore=12
    )
    normal_style = styles['Normal']
    
    # Título
    story.append(Paragraph("Informe de Detección YOLOv8 en Video", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Fecha de análisis: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Re sumen Ejecutivo
    story.append(Paragraph("Resumen Ejecutivo", heading_style))
    total_videos = len(all_detections)
    total_detections = sum(len(v['detections']) for v in all_detections)
    
    story.append(Paragraph(f"<b>Total de videos analizados:</b> {total_videos}", normal_style))
    story.append(Paragraph(f"<b>Total de objetos detectados:</b> {total_detections}", normal_style))
    
    if all_detections and 'modelo_usado' in all_detections[0]:
        story.append(Paragraph(f"<b>Modelo YOLO usado:</b> {all_detections[0]['modelo_usado']}", normal_style))
    
    if total_detections > 0:
        all_conf = [d['confidence'] for v in all_detections for d in v['detections']]
        avg_conf = float(np.mean(all_conf))
        avg_area = sum(d['area_pixels'] for v in all_detections for d in v['detections']) / total_detections
        unique_classes = set(d['class_name'] for v in all_detections for d in v['detections'])
        
        story.append(Paragraph(f"<b>Estado:</b> Se detectaron objetos en el/los video(s)", normal_style))
        story.append(Paragraph(f"<b>Confianza promedio:</b> {avg_conf:.3f}", normal_style))
        story.append(Paragraph(f"<b>Clases detect adas:</b> {', '.join(unique_classes)}", normal_style))
        story.append(Paragraph(f"<b>Área promedio detectada:</b> {avg_area:.0f} píxeles²", normal_style))
    else:
        story.append(Paragraph("<b>Estado:</b> No se detectaron objetos significativos", normal_style))
    
    story.append(Spacer(1, 20))
    
    # Análisis Detallado
    if total_detections > 0:
        story.append(PageBreak())
        story.append(Paragraph("Análisis Detallado por Video", heading_style))
        
        for i, video_data in enumerate(all_detections, 1):
            if not video_data['detections']:
                continue
            
            story.append(Paragraph(f"Video {i}: {video_data['video_name']}", heading_style))
            
            # Tabla de detecciones
            table_data = [['Frame', 'Tiempo (s)', 'Clase', 'Confianza', 'X', 'Y', 'Ancho', 'Alto', 'Área (%)']]
            
            for det in video_data['detections']:
                x, y, w, h = det['bbox']
                table_data.append([
                    str(det['frame']),
                    f"{det['time']:.2f}",
                    det['class_name'],
                    f"{det['confidence']:.3f}",
                    str(x),
                    str(y),
                    str(w),
                    str(h),
                    f"{det['area_percentage']:.2f}%"
                ])
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))
            story.append(table)
            story.append(Spacer(1, 20))
    
    # Metodología
    story.append(PageBreak())
    story.append(Paragraph("Metodología y Tecnología", heading_style))
    story.append(Paragraph("<b>Técnica de detección:</b> YOLOv8 (You Only Look Once versión 8)", normal_style))
    story.append(Paragraph("<b>Procesamiento:</b> Detección de objetos en tiempo real con bounding boxes y confianza", normal_style))
    story.append(Paragraph("<b>Umbral de confianza:</b> 0.5 (50%)", normal_style))
    story.append(Paragraph(f"<b>Dispositivo de procesamiento:</b> {device}", normal_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# ─── Informe de Continuidad - PDF ──────────────────────────────────────────────────────────────────────

def generate_continuity_report_pdf(analysis_results: List[Dict]) -> bytes:
    """
    Genera el informe PDF de análisis de continuidad.
    Retorna el documento como bytes.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    story = []
    styles = getSampleStyleSheet()
    temp_files = []  # Lista para rastrear archivos temporales
    
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a5490'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#1a5490'),
        spaceAfter=12,
        spaceBefore=12
    )
    normal_style = styles['Normal']
    
    # Título
    story.append(Paragraph("Informe de Análisis de Continuidad de Video", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Fecha de análisis: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Resumen Ejecutivo
    story.append(Paragraph("Resumen Ejecutivo", heading_style))
    total_videos = len(analysis_results)
    total_disc = sum(len(r['discontinuities']) for r in analysis_results)
    
    story.append(Paragraph(f"<b>Total de videos analizados:</b> {total_videos}", normal_style))
    story.append(Paragraph(f"<b>Total de discontinuidades detectadas:</b> {total_disc}", normal_style))
    
    if total_disc > 0:
        story.append(Paragraph("<b>Estado:</b> Se detectaron posibles cortes o ediciones", normal_style))
        story.append(Paragraph(f"<b>Promedio por video:</b> {total_disc / total_videos:.1f}", normal_style))
    else:
        story.append(Paragraph("<b>Estado:</b> No se detectaron discontinuidades significativas", normal_style))
    
    story.append(Spacer(1, 20))
    story.append(PageBreak())
    
    # Análisis Detallado
    for i, result in enumerate(analysis_results, 1):
        story.append(Paragraph(f"Análisis Detallado - Video {i}", heading_style))
        
        # Tabla de información
        info_data = [
            ['Nombre del archivo:', result['video_name']],
            ['Total de frames:', str(result['total_frames'])],
            ['FPS:', f"{result['fps']:.2f}"],
            ['Duración:', f"{result['duration']:.2f} segundos"],
            ['Distancia máxima:', f"{result['max_distance']:.4f}"],
            ['Discontinuidades detectadas:', str(len(result['discontinuities']))]
        ]
        
        info_table = Table(info_data, colWidths=[2.5*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Insertar gráfico si viene en base64
        if result.get('plot_base64'):
            story.append(Paragraph("Gráfico de Análisis de Continuidad", heading_style))
            try:
                plot_bytes = base64.b64decode(result['plot_base64'])
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(plot_bytes)
                    tmp_path = tmp.name
                temp_files.append(tmp_path)  # Guardar para eliminar después
                img = RLImage(tmp_path, width=6.5*inch, height=4*inch)
                story.append(img)
            except Exception as e:
                story.append(Paragraph(f"Error al insertar gráfico: {e}", normal_style))
            story.append(Spacer(1, 20))
        
        # Discontinuidades
        if result['discontinuities']:
            story.append(Paragraph("Discontinuidades Detectadas", heading_style))
            disc_data = [['Frame', 'Tiempo (s)', 'Tiempo (mm:ss)', 'Distancia']]
            
            for disc in result['discontinuities']:
                disc_data.append([
                    str(disc['frame']),
                    f"{disc['time']:.2f}",
                    disc.get('time_formatted', ''),
                    f"{disc['distance']:.4f}"
                ])
            
            disc_table = Table(disc_data)
            disc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(disc_table)
        else:
            story.append(Paragraph("No se detectaron discontinuidades significativas.", normal_style))
        
        if i < len(analysis_results):
            story.append(PageBreak())
    
    # Metodología
    story.append(PageBreak())
    story.append(Paragraph("Metodología Empleada", heading_style))
    story.append(Paragraph("<b>Técnica utilizada:</b> Análisis de correlación de histogramas", normal_style))
    story.append(Paragraph("<b>Umbral de detección:</b> Distancia &gt; 1.0", normal_style))
    story.append(Paragraph("<b>Procesamiento:</b> Conversión a escala de grises, histogramas normalizados, comparación por correlación entre frames consecutivos.", normal_style))
    
    # Construir el PDF
    doc.build(story)
    
    # Eliminar archivos temporales después de construir el PDF
    for tmp_path in temp_files:
        try:
            os.unlink(tmp_path)
        except:
            pass
    
    buffer.seek(0)
    return buffer.read()


# ─── Informe de Conversión a Escala de Grises - PDF ──────────────────────────────────────────────

def generate_grayscale_report_pdf(conversion_result: Dict) -> bytes:
    """
    Genera el informe PDF de conversión a escala de grises.
    Retorna el documento como bytes.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    story = []
    styles = getSampleStyleSheet()
    
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a5490'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#1a5490'),
        spaceAfter=12,
        spaceBefore=12
    )
    normal_style = styles['Normal']
    
    # Título
    story.append(Paragraph("Informe de Conversión a Escala de Grises", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Resumen de Conversión
    story.append(Paragraph("Resumen de Conversión", heading_style))
    
    total = conversion_result.get('total_files', 0)
    success = conversion_result.get('converted_count', 0)
    error_count = conversion_result.get('error_count', 0)
    rate = (success / total * 100) if total > 0 else 0
    
    summary_data = [
        ['Métrica', 'Valor'],
        ['Total de archivos', str(total)],
        ['Conversiones exitosas', str(success)],
        ['Errores', str(error_count)],
        ['Tasa de éxito', f'{rate:.1f}%']
    ]
    
    summary_table = Table(summary_data)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # Imágenes Convertidas
    if conversion_result.get('converted'):
        story.append(Paragraph("Imágenes Convertidas", heading_style))
        
        img_data = [['Original', 'Convertido', 'Resolución', 'Tamaño orig. (KB)', 'Reducción (%)']]
        
        for img in conversion_result['converted']:
            img_data.append([
                img['original_name'],
                img['output_name'],
                f"{img['original_size'][0]}x{img['original_size'][1]}",
                f"{img['original_file_size_kb']:.1f}",
                f"{img['size_reduction_pct']:.1f}%"
            ])
        
        img_table = Table(img_data)
        img_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(img_table)
        story.append(Spacer(1, 20))
    
    # Errores
    if conversion_result.get('errors'):
        story.append(Paragraph("Errores Encontrados", heading_style))
        for err in conversion_result['errors']:
            story.append(Paragraph(f"• {err}", normal_style))
        story.append(Spacer(1, 20))
    
    # Información Técnica
    story.append(PageBreak())
    story.append(Paragraph("Información Técnica", heading_style))
    
    tech_info = [
        ['Aspecto', 'Detalle'],
        ['Método', 'OpenCV cv2.cvtColor() con CV_BGR2GRAY'],
        ['Formato de salida', 'PNG para preservar calidad'],
        ['Espacio de color', 'Escala de grises (8 bits)']
    ]
    
    tech_table = Table(tech_info)
    tech_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(tech_table)
    
    # Construir el PDF
    doc.build(story)
    
    buffer.seek(0)
    return buffer.read()